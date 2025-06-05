import os
import docx
import pdfplumber
import chardet
import pytesseract
from pdf2image import convert_from_path
import tempfile
from tqdm import tqdm
from razdel import sentenize
import warnings
from typing import List, Dict, Optional, Union, Any
warnings.filterwarnings("ignore", message="CropBox missing from /Page, defaulting to MediaBox")

# в общем-то оно уникальное, надо его в PATH ещё добавить
TESSERACT_CMD = 'C:/Program Files/Tesseract-OCR/tesseract.exe'
POPPLER_PATH = "C:/Program Files/poppler-24.08.0/Library/bin"
pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD


def extract_text_from_docx(file_path):
    """Извлечение текста (включая таблицы) из файла DOCX."""
    try:
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))
        return "\n".join(full_text)
    except Exception as e:
        print(f"[ERROR] Error processing DOCX {file_path}: {e}")
        return ""

def extract_text_from_pdf(file_path, min_words_per_page=50):
    """
    Извлечение текста из PDF с помощью pdfplumber.
    Вычисляет количество страниц и среднее количество слов на странице. 
    Возвращается к OCR, если среднее значение слишком низкое.
    """
    text_pages = []
    page_word_counts = []
    pages_info = []
    try:
        with pdfplumber.open(file_path) as pdf:
            num_pages = len(pdf.pages)
            for i, page in enumerate(pdf.pages):
                try:
                    page_text = page.extract_text() or ""
                except Exception as e:
                    print(f"[ERROR] Text extraction failed on page {i+1} of {file_path}: {e}")
                    page_text = ""
                word_count = len(page_text.split())
                page_word_counts.append(word_count)
                pages_info.append({'page_number': i+1, 'word_count': word_count})
                text_pages.append(page_text)
        full_text = "\n".join(text_pages)
        avg_words = sum(page_word_counts) / len(page_word_counts) if page_word_counts else 0

        metadata = {
            'file_type': 'pdf',
            'num_pages': num_pages,
            'avg_words_per_page': avg_words,
            'pages_info': pages_info,
            'ocr_used': False
        }
        if avg_words < min_words_per_page:
            print(f"[INFO] Low average words per page ({avg_words:.2f} < {min_words_per_page}) for {file_path}. Switching to OCR.")
            full_text = extract_text_from_pdf_with_ocr(file_path)
            metadata['ocr_used'] = True
        
        return full_text, metadata
    except Exception as e:
        print(f"[ERROR] Error processing PDF {file_path}: {e}")
        return "", {}

def extract_text_from_pdf_with_ocr(file_path):
    """Извлечение текста из PDF с помощью OCR. 
    Пробует сначала сначала с poppler_path; если это не поможет, 
    То пробует без него.
    """
    try:
        print(f"[INFO] Starting OCR extraction for: {os.path.basename(file_path)}")
        all_text = []
        with tempfile.TemporaryDirectory() as temp_dir:
            try:
                images = convert_from_path(file_path, poppler_path=POPPLER_PATH, dpi=300)
            except Exception as e:
                print(f"[ERROR] During PDF conversion with poppler_path: {e}. Trying without poppler_path...")
                try:
                    images = convert_from_path(file_path, dpi=300)
                except Exception as e:
                    print(f"[ERROR] Conversion without poppler_path also failed: {e}")
                    return ""
            print(f"[INFO] Found {len(images)} pages in the PDF")
            for i, image in enumerate(tqdm(images, desc="Processing pages with OCR")):
                temp_file = os.path.join(temp_dir, f'page_{i}.png')
                image.save(temp_file, 'PNG')
                custom_config = r'--oem 3 --psm 6'
                try:
                    text = pytesseract.image_to_string(temp_file, lang='rus+eng', config=custom_config)
                    if not text.strip():
                        text = pytesseract.image_to_string(temp_file, lang='rus', config=custom_config)
                    if not text.strip():
                        text = pytesseract.image_to_string(temp_file, config=custom_config)
                except Exception as e:
                    print(f"[ERROR] OCR on page {i+1}: {e}")
                    text = pytesseract.image_to_string(temp_file, config=custom_config)
                all_text.append(text)
        full_text = "\n\n".join(all_text)
        return full_text
    except Exception as e:
        print(f"[ERROR] OCR processing failed for {file_path}: {e}")
        return ""

def extract_text_from_txt(file_path):
    try:
        with open(file_path, 'rb') as f:
            raw_data = f.read()
            detected_encoding = chardet.detect(raw_data)['encoding']
        with open(file_path, 'r', encoding=detected_encoding or 'utf-8', errors='replace') as file:
            text = file.read()
        metadata = {'file_type': 'txt'}
        return text, metadata
    except Exception as e:
        print(f"[ERROR] Error processing TXT {file_path}: {e}")
        return "", {}

def split_text_into_sentences(text):
    """Split text into sentences using razdel."""
    return [s.text for s in sentenize(text)]

def create_chunks_by_sentence(text, file_metadata, target_chunk_size=512, min_chunk_size=256, overlap_size=50):
    """
    Создавает текстовые фрагменты, объединяя предложения до тех пор, пока не будет достигнуто целевое количество слов,
    с перекрытием между фрагментами. Каждый фрагмент обогащается метаданными файла.
    """
    sentences = split_text_into_sentences(text)
    chunks = []
    current_chunk = []
    current_size = 0

    for i, sentence in enumerate(sentences):
        sentence_size = len(sentence.split())
        if current_size + sentence_size > target_chunk_size and current_chunk:
            chunk_text = " ".join(current_chunk)
            chunks.append(chunk_text)
            # overlap from the end of the current chunk
            if overlap_size > 0 and i > 0:
                overlap_sentences = []
                overlap_tokens = 0
                j = len(current_chunk) - 1
                while j >= 0 and overlap_tokens < overlap_size:
                    overlap_sentences.insert(0, current_chunk[j])
                    overlap_tokens += len(current_chunk[j].split())
                    j -= 1
                current_chunk = overlap_sentences + [sentence]
                current_size = sum(len(s.split()) for s in current_chunk)
            else:
                current_chunk = [sentence]
                current_size = sentence_size
        else:
            current_chunk.append(sentence)
            current_size += sentence_size

    if current_chunk:
        chunks.append(" ".join(current_chunk))

    # Build chunk dictionaries; metadata (including pages_info if present)
    chunk_dicts = []
    for idx, chunk in enumerate(chunks):
        if len(chunk.split()) >= min_chunk_size:
            data = {
                "id": f"{file_metadata.get('file_name', 'unknown')}-chunk-{idx}",
                "text": chunk,
                "metadata": {**file_metadata, "chunk_id": idx, "token_count": len(chunk.split())}
            }
            chunk_dicts.append(data)
    return chunk_dicts

def process_document(file_path, min_words_per_page=50, target_chunk_size=512, min_chunk_size=256, overlap_size=50):
    file_extension = os.path.splitext(file_path)[1].lower()
    file_metadata = {"file_name": os.path.basename(file_path), "file_path": file_path}
    text, meta = "", {}

    if file_extension == '.docx':
        text = extract_text_from_docx(file_path)
        meta = {"file_type": "docx"}
    elif file_extension == '.pdf':
        text, meta = extract_text_from_pdf(file_path, min_words_per_page)
    elif file_extension == '.txt':
        text, meta = extract_text_from_txt(file_path)
    else:
        print(f"[INFO] Unsupported file type: {file_path}")
        return []

    file_metadata.update(meta)
    # Always use the (possibly empty) text from OCR instead of skipping.
    # (Remove or adjust the below check if you want to include even very short content.)
    if len(text.strip()) < 50:
        print(f"[INFO] File {file_metadata['file_name']} produced less than 50 characters; proceeding anyway.")
    return create_chunks_by_sentence(text, file_metadata, target_chunk_size, min_chunk_size, overlap_size)

def process_document_folder(
    folder_path: str,
    min_words_per_page: int = 100,
    target_chunk_size: int = 512,
    min_chunk_size: int = 256,
    overlap_size: int = 150,
    include_metadata: bool = True  # Add this parameter
) -> List[Dict]:
    chunks = []
    for root, _, files in os.walk(folder_path):
        for file in files:
            if file.lower().endswith(('.pdf', '.docx', '.txt')):
                file_path = os.path.join(root, file)
                doc_chunks = process_document(
                    file_path,
                    min_words_per_page,
                    target_chunk_size,
                    min_chunk_size,
                    overlap_size
                )
                
                #Ensure metadata is added for each chunk
                for i, chunk in enumerate(doc_chunks):
                    if include_metadata:
                        chunk["metadata"] = {
                            "file_name": file,
                            "page": chunk.get("page", 0),
                            "source": file_path,
                            "chunk": i
                        }
                    chunks.extend(doc_chunks)
    return chunks
